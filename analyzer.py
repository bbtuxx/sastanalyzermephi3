import json
import sys
import urllib.parse
from datetime import datetime
from docx import Document
from odf.opendocument import OpenDocumentText
from odf.text import P
import zipfile
import os
import tempfile
import shutil
import re
import requests
import subprocess

GITHUB_BASE_URL = "https://raw.githubusercontent.com/bbtuxx/sastanalyzermephi3/main/"
REQUIREMENTS_URL = GITHUB_BASE_URL + "requirements.txt"

def install_requirements():
    try:
        import docx
        import odf
        import requests
    except ImportError:
        print("Необходимые библиотеки не установлены. Установка...")
        response = requests.get(REQUIREMENTS_URL)
        if response.status_code == 200:
            with open("requirements.txt", "wb") as f:
                f.write(response.content)
            subprocess.check_call([sys.executable, "-m", "pip", "install", "-r", "requirements.txt"])
        else:
            print(f"Не удалось загрузить requirements.txt с {REQUIREMENTS_URL}")
            sys.exit(1)

install_requirements()

def load_vulnerabilities(language):
    file_map = {
        "python": "vulnpython.json",
        "php": "vulnphp.json",
        "c": "vulnc.json",
        "csharp": "vulnsharp.json"
    }
    file_name = file_map.get(language)
    if file_name:
        url = GITHUB_BASE_URL + file_name
        print(f"Загрузка файла уязвимостей с {url}")
        response = requests.get(url)
        if response.status_code == 200:
            return response.json()["vulnerabilities"]
        else:
            print(f"Не удалось загрузить файл уязвимостей с {url}. Код ответа: {response.status_code}")
            sys.exit(1)
    else:
        print(f"Unsupported language: {language}")
        sys.exit(1)

def detect_language(file_path):
    extension_map = {
        ".py": "python",
        ".php": "php",
        ".c": "c",
        ".cpp": "c",
        ".cs": "csharp"
    }
    _, ext = os.path.splitext(file_path)
    return extension_map.get(ext)

def analyze_code(vulns, file_path):
    with open(file_path, 'r') as file:
        code_lines = file.readlines()

    report = []
    for i, line in enumerate(code_lines):
        for vuln in vulns:
            if re.search(vuln["pattern"], line):
                result = {
                    "name": vuln['name'],
                    "description": vuln['description'],
                    "recommendation": vuln['recommendation'],
                    "file": file_path,
                    "line_number": i + 1,
                    "code": line.strip(),
                    "search_url": f"https://www.google.com/search?q={urllib.parse.quote(vuln['search_query'])}"
                }
                report.append(result)
                print(f"[Уязвимость] {vuln['name']}: {vuln['description']}")
                print(f"Рекомендация: {vuln['recommendation']}")
                print(f"Файл: {file_path}, Строка: {i + 1}")
                print(f"Код: {line.strip()}")
                print(f"Подробнее: {result['search_url']}")
                print("-" * 80)

    return report

def save_report_to_word(report, output_file):
    doc = Document()
    doc.add_heading('Отчет об уязвимостях', 0)

    for item in report:
        doc.add_heading(item['name'], level=1)
        doc.add_paragraph(f"Описание: {item['description']}")
        doc.add_paragraph(f"Рекомендация: {item['recommendation']}")
        doc.add_paragraph(f"Файл: {item['file']}, Строка: {item['line_number']}")
        doc.add_paragraph(f"Код: {item['code']}")
        doc.add_paragraph(f"Подробнее: {item['search_url']}")
        doc.add_paragraph("-" * 80)

    doc.save(output_file)

def save_report_to_odt(report, output_file):
    textdoc = OpenDocumentText()
    for item in report:
        textdoc.text.addElement(P(text=f"[Уязвимость] {item['name']}: {item['description']}"))
        textdoc.text.addElement(P(text=f"Рекомендация: {item['recommendation']}"))
        textdoc.text.addElement(P(text=f"Файл: {item['file']}, Строка: {item['line_number']}"))
        textdoc.text.addElement(P(text=f"Код: {item['code']}"))
        textdoc.text.addElement(P(text=f"Подробнее: {item['search_url']}"))
        textdoc.text.addElement(P(text="-" * 80))

    textdoc.save(output_file)

def save_report_to_txt(report, output_file):
    with open(output_file, 'w') as file:
        for item in report:
            file.write(f"[Уязвимость] {item['name']}: {item['description']}\n")
            file.write(f"Рекомендация: {item['recommendation']}\n")
            file.write(f"Файл: {item['file']}, Строка: {item['line_number']}\n")
            file.write(f"Код: {item['code']}\n")
            file.write(f"Подробнее: {item['search_url']}\n")
            file.write("-" * 80 + "\n")

def create_zip_archive(files, output_zip):
    with zipfile.ZipFile(output_zip, 'w') as zipf:
        for file in files:
            zipf.write(file, arcname=os.path.basename(file))

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Использование: python analyzer.py <путь_к_проверяемому_файлу>")
        sys.exit(1)

    code_file = sys.argv[1]
    language = detect_language(code_file)
    if not language:
        print(f"Не поддерживаемый язык для файла: {code_file}")
        sys.exit(1)

    vulnerabilities = load_vulnerabilities(language)
    report = analyze_code(vulnerabilities, code_file)

    if report:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        temp_dir = tempfile.mkdtemp()
        word_file = os.path.join(temp_dir, f"vulnerability_report_{timestamp}.docx")
        odt_file = os.path.join(temp_dir, f"vulnerability_report_{timestamp}.odt")
        txt_file = os.path.join(temp_dir, f"vulnerability_report_{timestamp}.txt")
        save_report_to_word(report, word_file)
        save_report_to_odt(report, odt_file)
        save_report_to_txt(report, txt_file)

        zip_file = f"vulnerability_report_{timestamp}.zip"
        create_zip_archive([word_file, odt_file, txt_file], zip_file)

        shutil.rmtree(temp_dir)
        print(f"Отчет сохранен в архив {zip_file}.")
    else:
        print("Уязвимости не найдены.")
